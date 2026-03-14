from __future__ import annotations

import argparse
import random
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List

from docx import Document


@dataclass(frozen=True)
class RoleTemplate:
    title: str
    skills: List[str]
    project_focus: str


FIRST_NAMES = [
    "Aarav",
    "Vihaan",
    "Aditya",
    "Ishaan",
    "Kabir",
    "Rohan",
    "Arjun",
    "Nikhil",
    "Dev",
    "Karan",
    "Rahul",
    "Manav",
    "Ananya",
    "Aisha",
    "Riya",
    "Priya",
    "Sanya",
    "Neha",
    "Meera",
    "Diya",
    "Zoya",
    "Ira",
    "Sara",
    "Kavya",
    "Tara",
    "Aditi",
    "Siddharth",
    "Yash",
    "Harsh",
    "Nitin",
    "Bhavya",
    "Tanvi",
    "Pooja",
    "Sneha",
    "Varun",
    "Gaurav",
    "Mihir",
    "Pranav",
    "Shreya",
    "Naina",
]


LAST_NAMES = [
    "Sharma",
    "Verma",
    "Gupta",
    "Singh",
    "Patel",
    "Mehta",
    "Reddy",
    "Nair",
    "Iyer",
    "Kapoor",
    "Bansal",
    "Saxena",
    "Agarwal",
    "Kumar",
    "Mishra",
    "Joshi",
    "Malhotra",
    "Chopra",
    "Das",
    "Menon",
]


LOCATIONS = [
    "Bengaluru",
    "Hyderabad",
    "Pune",
    "Delhi",
    "Mumbai",
    "Chennai",
    "Noida",
    "Gurugram",
    "Kolkata",
    "Ahmedabad",
]


EDUCATION_LINES = [
    "B.Tech in Computer Science, IIT Delhi, 2016",
    "B.E. in Information Technology, NIT Trichy, 2015",
    "M.Tech in Data Science, IISc Bangalore, 2018",
    "B.Sc in Computer Science, Delhi University, 2014",
    "MCA, Pune University, 2017",
    "BCA, Christ University, 2016",
    "Master of Science in AI, IIIT Hyderabad, 2019",
    "Bachelor of Engineering in Software, Anna University, 2015",
]


CERTIFICATIONS = [
    "AWS Certified Solutions Architect",
    "Microsoft Certified: Azure AI Engineer Associate",
    "Google Professional Data Engineer",
    "TensorFlow Developer Certificate",
    "Databricks Certified Data Engineer",
    "Kubernetes and Cloud Native Associate",
    "HashiCorp Terraform Associate",
    "Oracle Java SE Professional",
]


ROLE_TEMPLATES = [
    RoleTemplate(
        title="Machine Learning Engineer",
        skills=["Python", "Machine Learning", "NLP", "PyTorch", "AWS", "Docker"],
        project_focus="NLP model deployment",
    ),
    RoleTemplate(
        title="Data Scientist",
        skills=["Python", "Machine Learning", "Pandas", "NumPy", "Scikit-learn", "SQL"],
        project_focus="predictive analytics",
    ),
    RoleTemplate(
        title="Backend Engineer",
        skills=["Python", "FastAPI", "PostgreSQL", "Redis", "Docker", "Git"],
        project_focus="high-scale API platform",
    ),
    RoleTemplate(
        title="Cloud Engineer",
        skills=["AWS", "Azure", "Terraform", "Kubernetes", "Linux", "Python"],
        project_focus="cloud migration",
    ),
    RoleTemplate(
        title="Full Stack Engineer",
        skills=["JavaScript", "TypeScript", "React", "Node.js", "SQL", "Docker"],
        project_focus="customer-facing web product",
    ),
    RoleTemplate(
        title="Java Engineer",
        skills=["Java", "Spring Boot", "SQL", "Docker", "Kubernetes", "Git"],
        project_focus="microservices modernization",
    ),
    RoleTemplate(
        title="Data Engineer",
        skills=["Python", "Spark", "Hadoop", "Airflow", "SQL", "AWS"],
        project_focus="batch and streaming pipelines",
    ),
    RoleTemplate(
        title="MLOps Engineer",
        skills=["Python", "Machine Learning", "Docker", "Kubernetes", "AWS", "Terraform"],
        project_focus="model serving and monitoring",
    ),
]


def sanitize_filename(name: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "_", name.strip().lower()).strip("_")
    return f"{slug}.docx"


def generate_candidate_name(index: int, rng: random.Random) -> str:
    first = FIRST_NAMES[index % len(FIRST_NAMES)]
    last = LAST_NAMES[(index * 3 + rng.randint(0, len(LAST_NAMES) - 1)) % len(LAST_NAMES)]
    return f"{first} {last}"


def years_to_dates(total_years: int) -> tuple[int, int]:
    current_year = datetime.now().year
    start_year = max(2008, current_year - total_years)
    midpoint = max(start_year + 1, current_year - max(1, total_years // 2))
    return start_year, midpoint


def write_resume(
    path: Path,
    name: str,
    role: RoleTemplate,
    years_exp: int,
    location: str,
    education_line: str,
    certification: str,
    rng: random.Random,
) -> None:
    start_year, midpoint = years_to_dates(years_exp)
    primary, secondary = role.skills[0], role.skills[1]
    tertiary = role.skills[2] if len(role.skills) > 2 else role.skills[0]
    skills_line = ", ".join(role.skills + rng.sample(["Git", "Linux", "SQL", "Docker", "Azure", "AWS"], k=2))
    skills_line = ", ".join(dict.fromkeys(skills_line.split(", ")))

    document = Document()
    document.add_heading(name, level=0)
    document.add_paragraph(f"{role.title} | {location}")
    document.add_paragraph(
        f"Email: {name.lower().replace(' ', '.')}@example.com | Phone: +91-98{rng.randint(10000000, 99999999)} | LinkedIn: linkedin.com/in/{name.lower().replace(' ', '-') }"
    )

    document.add_heading("Summary", level=1)
    document.add_paragraph(
        f"Results-driven {role.title} with {years_exp}+ years of experience in {primary} and {secondary}. "
        f"Delivered business impact through {role.project_focus}, scalable systems, and strong stakeholder collaboration."
    )

    document.add_heading("Skills", level=1)
    document.add_paragraph(skills_line)

    document.add_heading("Experience", level=1)
    document.add_paragraph(f"Senior {role.title} | NovaTech Systems | {midpoint} - Present")
    document.add_paragraph(
        f"- Led architecture and delivery of {role.project_focus} using {primary}, {secondary}, and {tertiary}."
    )
    document.add_paragraph(
        f"- Built production pipelines and improved performance by {rng.randint(25, 55)}% while mentoring a team of {rng.randint(2, 8)} engineers."
    )
    document.add_paragraph(f"{role.title} | PixelBridge Labs | {start_year} - {midpoint}")
    document.add_paragraph(
        f"- Implemented services with {primary}, {secondary}, and SQL; collaborated with product and analytics teams."
    )
    document.add_paragraph(
        f"- Achieved {years_exp}+ years hands-on experience in {primary} across enterprise-grade systems."
    )

    document.add_heading("Projects", level=1)
    document.add_paragraph(
        f"Intelligent Talent Platform: Designed and deployed a {role.project_focus} solution with {primary}, {secondary}, and Docker, supporting {rng.randint(50000, 200000)} monthly users."
    )
    document.add_paragraph(
        f"Observability and Reliability Upgrade: Introduced automated testing, CI/CD, and monitoring to reduce incidents by {rng.randint(20, 45)}%."
    )

    document.add_heading("Education", level=1)
    document.add_paragraph(education_line)

    document.add_heading("Certifications", level=1)
    document.add_paragraph(certification)

    document.save(path)


def generate_resumes(output_dir: Path, count: int, seed: int, overwrite: bool) -> List[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    rng = random.Random(seed)

    generated_files: List[Path] = []
    for idx in range(count):
        name = generate_candidate_name(idx, rng)
        role = ROLE_TEMPLATES[idx % len(ROLE_TEMPLATES)]
        years_exp = 2 + (idx % 10)
        location = LOCATIONS[idx % len(LOCATIONS)]
        education_line = EDUCATION_LINES[idx % len(EDUCATION_LINES)]
        certification = CERTIFICATIONS[idx % len(CERTIFICATIONS)]
        filename = sanitize_filename(name)
        file_path = output_dir / filename

        if file_path.exists() and not overwrite:
            generated_files.append(file_path)
            continue

        write_resume(
            path=file_path,
            name=name,
            role=role,
            years_exp=years_exp,
            location=location,
            education_line=education_line,
            certification=certification,
            rng=rng,
        )
        generated_files.append(file_path)

    return generated_files


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Generate synthetic sample resumes for the RAG assignment")
    parser.add_argument("--output-dir", type=str, default="resumes", help="Directory to write generated resumes")
    parser.add_argument("--count", type=int, default=35, help="Number of sample resumes to generate (>=30 recommended)")
    parser.add_argument("--seed", type=int, default=42, help="Random seed for deterministic output")
    parser.add_argument("--overwrite", action="store_true", help="Overwrite existing resume files")
    return parser


def main() -> None:
    parser = build_parser()
    args = parser.parse_args()

    if args.count < 30:
        parser.error("Please generate at least 30 resumes to satisfy assignment needs.")

    output_dir = Path(args.output_dir)
    files = generate_resumes(
        output_dir=output_dir,
        count=args.count,
        seed=args.seed,
        overwrite=args.overwrite,
    )

    print(f"Generated/available resumes: {len(files)}")
    print(f"Output directory: {output_dir.resolve().as_posix()}")


if __name__ == "__main__":
    main()
